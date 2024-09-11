# -- IMPORTS --

# -- LANGCHAIN IMPORTS --
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.chains import RetrievalQA
from langchain_community.llms import HuggingFaceEndpoint
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_pinecone import PineconeVectorStore
from langchain_community.document_loaders import S3DirectoryLoader
from langchain.chains import LLMChain
from langchain_core.prompts import PromptTemplate
from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate

# -- VECTOR DB IMPORTS --
from pinecone import Pinecone

# -- SERVER & HELPER IMPORTS --
from fastapi import FastAPI, BackgroundTasks
import requests
import concurrent.futures

# -- DOC READER - S3 IMPORTS --
import pandas as pd
import boto3
from io import BytesIO
from openpyxl import load_workbook
from docx import Document

# -- DATA CLEANING OPERATIONS IMPORTS --
import re
from sentence_transformers import SentenceTransformer, util
import torch

# -- CONFIG IMPORTS --
import os
from dotenv import load_dotenv


# -- CODE --

# nltk issue (uncomment)
#import nltk
#nltk.download('wordnet')

# load the environment variables
load_dotenv()


# -- FUNCTIONS --

# -- LOAD DOCUMENTS FROM S3 BUCKET --
def load_documents_from_s3_directory(directory_name, prefix):

    """Loads files (all types) from the specified S3 directory.

    Args:
        directory_name: The name of the S3 directory.
        prefix: Prefix of the specified S3 directory.

    Response:
        List of Loaded documents from the S3 directory.
    """


    document_loader = S3DirectoryLoader(
        directory_name, prefix = prefix)

    response = document_loader.load()

    return response

# -- TEXT SPLITTER HELPER FUNCTION --
def chunk_data(documents, chunk_size, chunk_overlap):

    """ Helper function that splits the incoming documents into chunks for further processing.
    
    Args: 
        documents (List): Raw documents loaded from the document loader.
        chunk_size (int): Size of a particular chunk for splitting 
        chunk_overlap (int): Overlap to maintain context between each chunk.
    """

    documents[0].page_content = re.sub(r'\s{3,}', ' ', documents[0].page_content)

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size,chunk_overlap=chunk_overlap)
    documents = text_splitter.split_documents(documents)

    return documents

# -- CREATING VECTOR STORE FROM DOCUMENTS AND INDEX NAME --
def create_vector_store_from_docs(documents, index_name):
   """Creates a vector store by processing documents with multithreading and adding chunks directly to the vector store."""
   
   Pinecone(api_key=os.environ['PINECONE_API_KEY'])
   
   model_name = "sentence-transformers/all-mpnet-base-v2"
   model_kwargs = {'device': 'cpu'}
   encode_kwargs = {'normalize_embeddings': False}
   embeddings = HuggingFaceEmbeddings(
       model_name=model_name,
       model_kwargs=model_kwargs,
       encode_kwargs=encode_kwargs
   )
   
   # Initialize the vector store
   vectorstore = PineconeVectorStore(index_name=index_name, embedding=embeddings)
   
   # Split documents into 4 parts for parallel processing
   num_docs = len(documents)
   split_docs = [documents[i::4] for i in range(4)]
   def worker(docs):
       """Worker function to process a chunk of documents and add to the vector store."""
       chunks = chunk_data(docs, 3500, 100)
       vectorstore.add_documents(chunks)
   
   # Execute the worker function in parallel using ThreadPoolExecutor
   with concurrent.futures.ThreadPoolExecutor(max_workers=4) as executor:
       executor.map(worker, split_docs)
   return vectorstore

# -- FETCH ALL THE QUESTIONS FROM THE RAW TEXT USING LLM CALL --
def fetch_questions_from_document(document):

    """For the given document in the argument, retrieve the questions and append to a string / list
 
    Args:
        document (list) : The document from which the questions needs to be extracted
 
    Return:
        list: A list of questions
 
    """
    file_data_string = document

    repo_id = "meta-llama/Meta-Llama-3.1-8B-Instruct"

    llm = HuggingFaceEndpoint(
        repo_id=repo_id,
        # max_new_tokens=4096,
        temperature=0.2,
        huggingfacehub_api_token=os.environ['HUGGINGFACEHUB_API_TOKEN'],
    )

    instruction = "Retrieve all the questions from the text given in the Document while making sure not to generate any duplicates. If there is follow-up questions to a given question, make sure it is included in the same line and not with line breaks. Make sure while returning the question, each question has two linebreaks between."
    context = file_data_string

    template = """Instruction: {instruction}
 
    Document: {context}
 
    Answer: Give out only the list of questions and nothing else."""

    prompt = template.format(instruction=instruction, context=context)
    file_question_content = llm(prompt)

    return file_question_content

# -- REMOVAL OF SEMANTIC DUPLICATE SENTENCES
def remove_semantic_duplicates(sentences: list):

    model = SentenceTransformer("all-MiniLM-L6-v2")
    paraphrases = util.paraphrase_mining(model, sentences)

    uniques = []
    for paraphrase in paraphrases:
        score, i, j = paraphrase
        #print("{} \t\t {} \t\t Score: {:.4f}".format(sentences[i], sentences[j], score))
        if(score > 0.8):
            uniques.append(sentences[i])

    for i in range(len(sentences) - 1, -1, -1):
        if sentences[i] in uniques:
            sentences.pop(i)

    return sentences

# -- FETCHING RESULTS FROM VECTORDB --
def fetch_result_from_context(question_list, vector_store):
    """For all the questions present in the document, fetch answers from the context and append to a list.
 
    Args:
        question_list (list): A list of questions.
        vector_store: The vector store.
 
    Returns:
        tuple: A tuple containing a big string with questions followed by answers and a list of individual question-answer pairs.
    """

    max_workers = 4  # Adjust this value based on your system and workload

    def process_questions(questions_chunk):
        results = []
        for question in questions_chunk:
            query = f"{question}"
            result = f"Question: {question}\n"
            result += new_retrieval_chain(vector_store, query)
            result += '\n\n'
            results.append(result)
        return results

    # Determine the size of each chunk for multithreading
    chunk_size = len(question_list) // max_workers
    if chunk_size == 0:
        chunk_size = 1
    
    chunks = [question_list[i:i + chunk_size] for i in range(0, len(question_list), chunk_size)]

    combined_results = []
    results_list = []

    # Use ThreadPoolExecutor to process questions in parallel
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = [executor.submit(process_questions, chunk) for chunk in chunks]
        for future in concurrent.futures.as_completed(futures):
            chunk_results = future.result()
            results_list.extend(chunk_results)
            combined_results.append(''.join(chunk_results))

    # Combine all results into a single string
    combined_results_string = ''.join(combined_results)

    return combined_results_string, results_list

# -- CREATE RETRIEVAL CHAIN SETUP --
def new_retrieval_chain(vector_database, input_query: str):

    # repo_id = "meta-llama/Meta-Llama-3.1-8B-Instruct"
    repo_id = "mistralai/Mistral-7B-Instruct-v0.2"

    large_language_model = HuggingFaceEndpoint(
        repo_id=repo_id, temperature=0.3, token=os.environ['HUGGINGFACEHUB_API_TOKEN']
    )


    system_prompt = (
        "You are an assistant that does not deviate from the instructions for answering a question. "
        "Use the following pieces of retrieved context to answer "
        "the question mentioned in the human prompt. If you don't know the answer, say that you "
        "don't know. Use three sentences maximum and keep the "
        "answer concise."
        "\n\n"
        "It is very important that you only answer to the question given by the human and you should never make up your own question"
        "The structure of your answer should always be:"
        "Answer: answer to the question"
        "{context}"
    )

    prompt = ChatPromptTemplate.from_messages(
        [
            ("system", system_prompt),
            ("human", "{input}"),
        ]
    )

    retriever = vector_database.as_retriever()

    question_answer_chain = create_stuff_documents_chain(large_language_model, prompt)
    rag_chain = create_retrieval_chain(retriever, question_answer_chain)

    response = rag_chain.invoke({"input": input_query})
    return response["answer"]

# -- WRITING RESPONSE OF QUESTIONNAIRE BACK TO S3 BUCKET --
def write_to_s3_folder(bucket_name, folder_name, file_name, content):
  """Writes content to a specific folder in an S3 bucket using put_object.

  Args:
    bucket_name: The name of the S3 bucket.
    folder_name: The name of the folder in the bucket.
    file_name: The name of the file.
    content: The content to be written to the file.
  """

  s3 = boto3.client('s3')
  object_key = f"{folder_name}/{file_name}"

  try:
    s3.put_object(Body=content, Bucket=bucket_name, Key=object_key)
    print(f"File uploaded successfully to S3: s3://{bucket_name}/{object_key}")
  except Exception as e:
    print(f"Error uploading file to S3: {e}")


# -- EXCEL WRITING FUNCTIONS --
def find_matching_columns(df, keywords):
    """Find columns that contain any of the specified keywords."""
    matching_columns = []
    for col in df.columns:
        if any(keyword.lower() in col.lower() for keyword in keywords):
            matching_columns.append(col)
    return matching_columns

def find_next_empty_cell(row):
    """Find the index of the next empty cell in a given row."""
    for col_idx, cell in enumerate(row):
        if cell.value is None or cell.value == "":
            return col_idx
    return None  # Return None if no empty cell is found

def extract_answer(string):
    """Extract the answer part of the string after 'Answer:'."""
    if "Answer:" in string:
        return string.split("Answer:")[1].strip()
    return None

# def update_answer_column(ws, df, query_column, strings_list):
#     """Update the next empty cell in a row based on matches in the query column."""
#     query_col_idx = df.columns.get_loc(query_column) + 1  # Adjust for 1-based index in openpyxl
#     for idx, query in enumerate(df[query_column]):
#         if isinstance(query, str):  # Check if the query is a string
#             for string in strings_list:
#                 if query.strip().lower() in string.lower():
#                     next_empty_col_idx = find_next_empty_cell(ws[idx + 2])  # Adjust for header row
#                     if next_empty_col_idx is not None:
#                         answer = extract_answer(string)  # Extract only the answer part
#                         if answer:
#                             ws.cell(row=idx + 2, column=next_empty_col_idx + 1, value=answer)  # Write the answer to the next empty cell
#                     break  # Move to the next row after finding a match

def update_answer_column(ws, df, query_column, strings_list, threshold=0.7):
    """Update the next empty cell in a row based on semantic similarity matches in the query column."""
    model = SentenceTransformer("all-MiniLM-L6-v2")
    
    query_col_idx = df.columns.get_loc(query_column) + 1  # Adjust for 1-based index in openpyxl
    
    # Encode all the queries and strings for comparison
    query_embeddings = model.encode(df[query_column].astype(str).tolist(), convert_to_tensor=True)
    string_embeddings = model.encode(strings_list, convert_to_tensor=True)
    
    for idx, query in enumerate(df[query_column]):
        if isinstance(query, str):
            # Compute cosine similarities between the current query and all strings
            similarities = util.cos_sim(query_embeddings[idx], string_embeddings)
            max_similarity, max_index = torch.max(similarities, dim=1)
            
            if max_similarity.item() >= threshold:
                matching_string = strings_list[max_index.item()]
                next_empty_col_idx = find_next_empty_cell(ws[idx + 2])  # Adjust for header row
                if next_empty_col_idx is not None:
                    answer = extract_answer(matching_string)  # Extract only the answer part
                    if answer:
                        ws.cell(row=idx + 2, column=next_empty_col_idx + 1, value=answer)  # Write the answer to the next empty cell

def process_excel_file(idx, strings_list, bucket_name, prefix):

    # Initialize S3 client
    s3 = boto3.client('s3')

    # # S3 bucket and file information
    # bucket_name = 'your-bucket-name'
    # prefix = 'path/to/your/folder/'  

    # List the files in the S3 bucket
    response = s3.list_objects_v2(Bucket=bucket_name, Prefix=prefix)

    # Filter the files to only include Excel files
    excel_files = [item['Key'] for item in response.get('Contents', []) if item['Key'].endswith('.xlsx')]

    file_key = excel_files[idx]
    print(file_key)
    # Download the file to a BytesIO object
    file_stream = BytesIO()
    s3.download_fileobj(bucket_name, file_key, file_stream)

    # Seek to the beginning of the BytesIO object
    file_stream.seek(0)

    # Define the keywords to search for in the columns
    keywords = ["query", "questions", "checklist", "description"]

    # Load the Excel file with openpyxl
    wb = load_workbook(file_stream)
    sheet_names = wb.sheetnames

    # Iterate through each sheet
    for sheet_name in sheet_names:
        ws = wb[sheet_name]
        df = pd.read_excel(file_stream, sheet_name=sheet_name)
        
        # Identify columns that contain the keywords
        matching_columns = find_matching_columns(df, keywords)

        # Update the next empty cell in the row for each matching column
        for column in matching_columns:
            update_answer_column(ws, df, column, strings_list)

    print('wrote to excel')

    # Save the updated workbook back to the Excel file
    save_to_s3(wb, bucket_name, file_key)

def save_to_s3(workbook, bucket_name, file_key):
    # Save the workbook to a BytesIO object
    output_stream = BytesIO()
    workbook.save(output_stream)
    output_stream.seek(0)

    # object_key = f"{folder_name}/{file_name}"
    file_key_list = file_key.split('/')
    file_key_list[1] = 'result'
    file_key_str = '/'.join(file_key_list)
    print(file_key_str)
    # Upload the file back to S3
    s3 = boto3.client('s3')
    s3.put_object(Bucket=bucket_name, Key=file_key_str, Body=output_stream)
    print('wrote to s3')

# -- DOCX WRITING FUNCTIONS --

def update_document_with_answers(doc, sentences, strings_list, threshold=0.7):
    """Update the document by appending answers to sentences that match with strings in the result list."""
    model = SentenceTransformer("all-MiniLM-L6-v2")
    
    # Encode all the sentences and strings for comparison
    sentence_embeddings = model.encode(sentences, convert_to_tensor=True)
    string_embeddings = model.encode(strings_list, convert_to_tensor=True)
    
    for idx, sentence in enumerate(sentences):
        # Compute cosine similarities between the current sentence and all strings
        similarities = util.cos_sim(sentence_embeddings[idx], string_embeddings)
        max_similarity, max_index = torch.max(similarities, dim=1)
        
        if max_similarity.item() >= threshold:
            matching_string = strings_list[max_index.item()]
            answer = extract_answer(matching_string)  # Extract only the answer part
            if answer:
                # Append the answer in a new line after the sentence
                doc.paragraphs[idx].add_run(f'\nAnswer: {answer}')

def process_word_file(idx, strings_list, bucket_name, prefix):

    # Initialize S3 client
    s3 = boto3.client('s3')

    # List the files in the S3 bucket
    response = s3.list_objects_v2(Bucket=bucket_name, Prefix=prefix)

    # Filter the files to only include Word documents
    word_files = [item['Key'] for item in response.get('Contents', []) if item['Key'].endswith('.docx')]

    file_key = word_files[idx]
    print(file_key)
    
    # Download the file to a BytesIO object
    file_stream = BytesIO()
    s3.download_fileobj(bucket_name, file_key, file_stream)

    # Seek to the beginning of the BytesIO object
    file_stream.seek(0)

    # Load the Word document
    doc = Document(file_stream)

    # Extract all sentences from the document
    sentences = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]

    # Update the document with answers
    update_document_with_answers(doc, sentences, strings_list)

    print('Updated Word document')

    # Save the updated document back to the Word file in S3
    save_word_to_s3(doc, bucket_name, file_key)

def save_word_to_s3(document, bucket_name, file_key):
    # Save the document to a BytesIO object
    output_stream = BytesIO()
    document.save(output_stream)
    output_stream.seek(0)

    # Modify the file key to indicate the result folder
    file_key_list = file_key.split('/')
    file_key_list[1] = 'result'
    file_key_str = '/'.join(file_key_list)
    print(file_key_str)
    
    # Upload the file back to S3
    s3 = boto3.client('s3')
    s3.put_object(Bucket=bucket_name, Key=file_key_str, Body=output_stream)
    print('Uploaded to S3')


async def retrieval_llama(user_id, req_id):
    #import nltk
    #nltk.download()
    #return {"message":"success"}
    # # -- CREATING VECTOR STORE --
    print('call started')
    bucket_name = 'ofofo-stage-storage'
    prefix = 'questionnaire/context/' + user_id + '/' + req_id + '/'
    context_file_data = load_documents_from_s3_directory(bucket_name, prefix)
    print("reading context done..!")
    # # Split the documents into chunks and create a vector store from the chunks
    documents = chunk_data(context_file_data, 3500, 100)
    index_name = 'docs-rag-chatbot'
    vector_store = create_vector_store_from_docs(documents, index_name)
    print("vector db created..!")

    embedding_function = HuggingFaceEmbeddings()

    index_name = "docs-rag-chatbot"

    vector_store = PineconeVectorStore(
        index_name = index_name, embedding=embedding_function
    )

    # -- FETCHING QUESTIONS FROM DOCS --
    # bucket_name = 'prestage-ofofo-storage'
    bucket_name = 'ofofo-stage-storage'
    #prefix = 'questions/'
    prefix = 'questionnaire/query/' + user_id + '/' + req_id + '/'
    questionnaires = load_documents_from_s3_directory(bucket_name, prefix)
    print("reading questionnaire done..!")
    print("LEN :::", len(questionnaires))
    # with open("a.txt", 'w') as f:
    #     f.write(str(questionnaires))
    #print(questionnaires[0])
    # -- FETCH RESULTS -- 
    idx = 0
    for questionnaire in questionnaires:

        chunks = chunk_data([questionnaire], 3500, 10)
        print('chunks created!')
        print("LEN CHUNKS :::", len(chunks))

        master_question_list = []

        for chunk in chunks:
            # print(chunks[0].page_content)
            # print(chunks[0])
            questions = ''
            questions = fetch_questions_from_document(chunk.page_content)
            questions += "\n\n"
            questions_list = [question.strip() for question in questions.strip().split('\n\n')]
            questions_list = remove_semantic_duplicates(questions_list)

            for question in questions_list:
                master_question_list.append(question)

        # for question in master_question_list:
        #     with open("ques.txt", 'a', encoding="utf-8") as f:
        #         f.write(question)
        #         f.write('\n\n')

        result_folder_name = 'questionnaire/result/' + user_id + '/' + req_id
        result_file_name = questionnaire.metadata["source"].split("/")[-1].split(".")[0]
        
        result, result_list = fetch_result_from_context(master_question_list, vector_store)
        print("results fetched::::::")

       # with open(result_file_name, 'a', encoding="utf-8") as f:
        #     f.write(result)
        #     f.write('\n')

        # need to write to excel before pushing it to s3
        # process_excel_file(file_path, strings_list)
        
        # bucket_name = 'prestage-ofofo-storage'
        bucket_name = 'ofofo-stage-storage'
        #prefix = 'questions/'
        prefix = 'questionnaire/query/' + user_id + '/' + req_id + '/'

        # write_to_s3_folder(bucket_name, result_folder_name, result_file_name, result)
        process_excel_file(idx, result_list, bucket_name, prefix)
        process_word_file(idx, result_list, bucket_name, prefix)
        
        print("results written to s3..!")
        print(idx, 'Index name')
        idx += 1

    print("everything is done ..!")
    # change
    questionnaire_status_url = "https://buyer-api.ofofo.io/api/v1/aether-free-api/questionnaire/" + user_id + "/" + req_id + "/status"
    print(questionnaire_status_url)

    requests.get(questionnaire_status_url)
    print("status updated")
    return {"message":"successsss"}

async def retrieval_llama_background(background_tasks, user_id, req_id):
    background_tasks.add_task(retrieval_llama, user_id, req_id)






